from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import pytesseract
from PIL import Image, ImageEnhance
from pdf2image import convert_from_path
import os
import shutil
import gc
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = FastAPI()

# تنظیمات هدر برای دسترسی بدون محدودیت مرورگر و حل خطای Mixed Content
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# تابع تنظیم راست‌چین کردن پاراگراف‌ها در فایل ورد برای زبان فارسی
def set_paragraph_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# تابع پردازش تصویر اصلی شما برای حفظ بالاترین کیفیت استخراج[span_1](start_span)[span_1](end_span)
def do_ocr_on_one_page(image_file_path, page_number):
    def improve_image_quality(img):
        gray = img.convert('L')
        enhancer = ImageEnhance.Contrast(gray)
        enhanced = enhancer.enhance(2.5) # کنتراست اصلی و تست‌شده شما[span_2](start_span)[span_2](end_span)
        return enhanced

    def make_sharp(img):
        sharpener = ImageEnhance.Sharpness(img)
        sharp_img = sharpener.enhance(1.3) # شارپنس اصلی و تست‌شده شما[span_3](start_span)[span_3](end_span)
        return sharp_img

    try:
        image = Image.open(image_file_path)
        better_image = improve_image_quality(image)
        final_image = make_sharp(better_image)
        
        # استخراج متن با کانفیگ اصلی جهت جلوگیری از درهم‌ریختگی کلمات[span_4](start_span)[span_4](end_span)
        extracted_text = pytesseract.image_to_string(final_image, lang='fas', config='--psm 3 --oem 3')
        
        image.close()
        better_image = None
        final_image = None
        del image, better_image, final_image
        
        try: os.remove(image_file_path)
        except: pass
            
        return {
            'page_num': page_number,
            'text_content': extracted_text.strip(),
            'character_count': len(extracted_text.strip()),
            'is_successful': True
        }
    except Exception as error:
        try: os.remove(image_file_path)
        except: pass
        return {
            'page_num': page_number,
            'text_content': f"خطا در صفحه {page_number}: {str(error)}",
            'character_count': 0,
            'is_successful': False
        }

@app.get("/")
def read_root():
    return {"status": "سرور ابری خروجی ورد فعال است"}

@app.post("/process-pdf")
async def process_pdf(file: UploadFile = File(...)):
    temp_pdf = "input.pdf"
    output_docx = "OCR_Result.docx"
    temp_dir = "/tmp/temp_ocr_images"
    
    # ذخیره فایل دریافتی از فرانت‌آند
    with open(temp_pdf, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    if os.path.exists(temp_dir): shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    # تعیین داینامیک DPI بر اساس حجم فایل ورودی برای بهینه‌سازی سرعت و رم[span_5](start_span)[span_5](end_span)
    pdf_size_mb = os.path.getsize(temp_pdf) / (1024 * 1024)
    if pdf_size_mb < 5:
        image_quality_dpi = 400
    elif pdf_size_mb < 20:
        image_quality_dpi = 350
    else:
        image_quality_dpi = 300
        
    all_ocr_results = []
    page_num = 1
    
    # پردازش ترتیبی و جریان‌محور صفحات برای جلوگیری از کرش سرور ابری
    while True:
        try:
            images = convert_from_path(temp_pdf, dpi=image_quality_dpi, first_page=page_num, last_page=page_num, fmt='jpeg')
            if not images:
                break
                
            file_name = f"page_{page_num:04d}.jpg"
            full_path = os.path.join(temp_dir, file_name)
            images[0].save(full_path, 'JPEG', quality=95, optimize=True)
            
            del images
            gc.collect()
            
            page_result = do_ocr_on_one_page(full_path, page_num)
            all_ocr_results.append(page_result)
            page_num += 1
        except:
            break

    # مرتب‌سازی صفحات بر اساس شماره صفحه برای تضمین عدم جابه‌جایی متون خروجی[span_6](start_span)[span_6](end_span)
    all_ocr_results.sort(key=lambda x: x['page_num'])
    
    # ساخت سند ورد مایکروسافت (docx)
    doc = Document()
    
    # تنظیم استایل و ابعاد فونت پیش‌فرض برای متون فارسی
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(13)
    
    # نوشتن اطلاعات آماری سند در ابتدای فایل ورد[span_7](start_span)[span_7](end_span)
    total_chars = sum(res['character_count'] for res in all_ocr_results if res['is_successful'])
    successful_count = sum(1 for res in all_ocr_results if res['is_successful'])
    
    p_title = doc.add_paragraph("نتایج سیستم هوشمند OCR فارسی ابری")
    set_paragraph_rtl(p_title)
    p_title.runs[0].font.bold = True
    p_title.runs[0].font.size = Pt(16)
    
    p_meta = doc.add_paragraph(f"فایل مبدا: {file.filename}\nتاریخ پردازش: {time.strftime('%Y-%m-%d %H:%M:%S')}\nصفحات موفق: {successful_count} از {len(all_ocr_results)}")
    set_paragraph_rtl(p_meta)
    
    doc.add_paragraph("=" * 60)
    
    # اضافه کردن متون استخراج شده به ساختار فایل ورد
    for single_result in all_ocr_results:
        p_page = doc.add_paragraph(f"\n--- صفحه {single_result['page_num']} ---")
        set_paragraph_rtl(p_page)
        p_page.runs[0].font.bold = True
        
        p_text = doc.add_paragraph(single_result['text_content'])
        set_paragraph_rtl(p_text)

    # ذخیره نهایی سند در فضای موقت سرور
    doc.save(output_docx)

    # پاکسازی فایل‌های زائد از روی هارد سرور
    if os.path.exists(temp_dir): shutil.rmtree(temp_dir)
    if os.path.exists(temp_pdf): os.remove(temp_pdf)
    
    # ارسال فایل ورد نهایی به مرورگر با هدرهای باز
    response = FileResponse(
        output_docx, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        filename="OCR_Result.docx"
    )
    response.headers["Access-Control-Allow-Origin"] = "*"
    return response
